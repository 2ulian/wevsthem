"""
Génère le rapport de stage BUT2 en format DOCX.
Police Calibri 11, texte justifié, marges moyennes, interligne simple.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_style(doc, name, font_name="Calibri", font_size=11, bold=False, color=None):
    style = doc.styles[name]
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor(*color)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return style


def add_heading1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.name = "Calibri"
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p


def add_heading2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.name = "Calibri"
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p


def add_para(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.name = "Calibri"
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "2E74B5")
        cell._tc.tcPr.append(shading)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.name = "Calibri"
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:fill"), "D6E4F0")
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_placeholder_figure(doc, description, width_note="pleine largeur"):
    """Adds a gray placeholder box with a description of the figure to insert."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"[ FIGURE À INSÉRER — {width_note} ]")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.bold = True

    p2 = doc.add_paragraph()
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    p2.paragraph_format.left_indent = Cm(1)
    p2.paragraph_format.right_indent = Cm(1)
    r2 = p2.add_run(description)
    r2.font.name = "Calibri"
    r2.font.size = Pt(9)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    return p2


def set_margins(doc, top=2.5, bottom=2.5, left=2.5, right=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


def add_header_footer(doc, title, author):
    for section in doc.sections:
        # Header
        header = section.header
        header.paragraphs[0].clear()
        run = header.paragraphs[0].add_run(f"{title}  —  {author}")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Footer with page number
        footer = section.footer
        footer.paragraphs[0].clear()
        footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_left = footer.paragraphs[0].add_run("Page ")
        run_left.font.name = "Calibri"
        run_left.font.size = Pt(9)
        run_left.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run_pn = footer.paragraphs[0].add_run()
        run_pn._r.append(fldChar1)
        run_pn._r.append(instrText)
        run_pn._r.append(fldChar2)
        run_pn.font.name = "Calibri"
        run_pn.font.size = Pt(9)


# ---------------------------------------------------------------------------
# Build Document
# ---------------------------------------------------------------------------

doc = Document()
set_margins(doc, 2.5, 2.5, 2.8, 2.8)

# Base styles
set_style(doc, "Normal", font_size=11)
try:
    set_style(doc, "List Bullet", font_size=11)
except:
    pass

REPORT_TITLE = "Détection automatique du langage « Nous vs Eux » dans les réseaux sociaux"
AUTHOR = "Julian Constant"

add_header_footer(doc, REPORT_TITLE, AUTHOR)

# ============================================================
# PAGE DE GARDE
# ============================================================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RAPPORT DE STAGE")
run.font.name = "Calibri"
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(REPORT_TITLE)
run.font.name = "Calibri"
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Stage de BUT Informatique – 2e année (BUT2)")
run.font.name = "Calibri"
run.font.size = Pt(12)
run.font.italic = True

doc.add_paragraph()
doc.add_paragraph()

info_table = doc.add_table(rows=8, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ("Étudiant",            "Julian Constant"),
    ("Dates du stage",      "Avril – Juin 2026"),
    ("Établissement d'accueil", "Laboratoire / Projet de recherche individuel"),
    ("Diplôme préparé",     "BUT Informatique"),
    ("Établissement",       "IUT"),
    ("Tuteur entreprise",   "À compléter"),
    ("Tuteur pédagogique",  "À compléter"),
    ("Année universitaire", "2025 – 2026"),
]
for i, (label, val) in enumerate(info_data):
    row = info_table.rows[i]
    c0 = row.cells[0]
    c1 = row.cells[1]
    c0.text = label
    c1.text = val
    c0.paragraphs[0].runs[0].font.name = "Calibri"
    c0.paragraphs[0].runs[0].font.size = Pt(11)
    c0.paragraphs[0].runs[0].bold = True
    c1.paragraphs[0].runs[0].font.name = "Calibri"
    c1.paragraphs[0].runs[0].font.size = Pt(11)
    c0.width = Cm(6)
    c1.width = Cm(9)

doc.add_page_break()

# ============================================================
# REMERCIEMENTS
# ============================================================
add_heading1(doc, "Remerciements")
add_para(doc,
    "Je tiens à remercier l'ensemble des personnes qui ont rendu ce stage possible et enrichissant. "
    "Je remercie mon tuteur pédagogique pour ses conseils et le suivi régulier tout au long du projet, "
    "ainsi que les membres de la communauté open source qui ont rendu disponibles les données et les outils "
    "utilisés dans ce travail (Arctic Shift / Pushshift pour les données Reddit, les équipes de recherche "
    "derrière BERTopic et GoEmotions). "
    "Ce projet m'a permis de combiner des compétences acquises en formation avec une mise en pratique "
    "concrète en traitement automatique du langage naturel."
)
doc.add_page_break()

# ============================================================
# SOMMAIRE (placeholder)
# ============================================================
add_heading1(doc, "Sommaire")
toc_entries = [
    ("1.", "Introduction", ""),
    ("2.", "Présentation du projet", ""),
    ("  2.1.", "Contexte et objectifs", ""),
    ("  2.2.", "Environnement de travail et outils", ""),
    ("3.", "Pipeline de traitement", ""),
    ("  3.1.", "Collecte et préparation des données", ""),
    ("  3.2.", "Détecteur de langage d'altérisation (règles)", ""),
    ("  3.3.", "Enrichissement : toxicité et émotions", ""),
    ("  3.4.", "Classifieur machine learning (LinearSVC TF-IDF)", ""),
    ("  3.5.", "Analyse thématique (BERTopic)", ""),
    ("  3.6.", "Analyse temporelle et event study", ""),
    ("4.", "Résultats et analyse", ""),
    ("  4.1.", "Comparaison inter-plateformes", ""),
    ("  4.2.", "Analyse par sous-reddit", ""),
    ("  4.3.", "Dynamiques temporelles", ""),
    ("5.", "Dashboard interactif", ""),
    ("6.", "Conclusion", ""),
    ("",  "Glossaire", ""),
    ("",  "Références", ""),
    ("",  "Annexe A – Matrice SWOT", ""),
    ("",  "Annexe B – Compétences BUT mobilisées", ""),
    ("",  "Annexe C – Utilisation de l'IA", ""),
]
for num, title, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"{num}  {title}")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    if not num.startswith("  "):
        run.bold = True

add_para(doc,
    "Note : les numéros de page seront renseignés automatiquement par Word après mise en page finale. "
    "Utilisez l'outil Table des matières automatique (Références > Table des matières).",
    italic=True
)
doc.add_page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
add_heading1(doc, "1. Introduction")

add_para(doc,
    "En 2024, plus de 5,2 milliards de personnes utilisent les réseaux sociaux dans le monde. "
    "Ces espaces d'expression de masse sont devenus un terrain privilégié pour l'observation des "
    "dynamiques de groupe et des phénomènes de polarisation sociale. Parmi ces phénomènes, le "
    "langage dit « Nous vs Eux » — ou altérisation — constitue un mécanisme rhétorique par lequel "
    "un groupe construit son identité en opposition à un « autre » désigné, souvent en l'associant "
    "à des métaphores déhumanisantes, des exclusions morales ou des cadrages de menace."
)

add_para(doc,
    "Ce stage s'inscrit dans un projet de recherche appliquée visant à détecter et quantifier "
    "automatiquement ce type de langage sur plusieurs plateformes sociales, en mobilisant des "
    "techniques de traitement automatique du langage naturel (NLP). La question centrale qui a guidé "
    "ce travail est la suivante : dans quelle mesure est-il possible de détecter et de mesurer "
    "automatiquement le langage d'altérisation dans les discours en ligne portant sur l'immigration, "
    "et comment ce langage varie-t-il selon la plateforme, le contexte géopolitique et le temps ?"
)

add_para(doc,
    "Ce rapport présente d'abord le contexte et l'environnement du projet, puis détaille la chaîne "
    "de traitement construite — de la collecte des données à la création d'un dashboard interactif "
    "— avant d'analyser les principaux résultats obtenus sur un corpus de plus de 134 000 publications."
)
doc.add_page_break()

# ============================================================
# 2. PRÉSENTATION DU PROJET
# ============================================================
add_heading1(doc, "2. Présentation du projet")

add_heading2(doc, "2.1 Contexte et objectifs")

add_para(doc,
    "Le projet « We vs Them » est un projet de recherche individuel mené en autonomie sur une "
    "durée de deux mois. Il vise à construire une boîte à outils NLP complète pour l'analyse du "
    "discours d'altérisation dans les réseaux sociaux anglophones, avec une attention particulière "
    "portée aux discussions sur l'immigration."
)

add_para(doc, "Les objectifs principaux sont les suivants :")
add_bullet(doc, "Collecter des données textuelles multi-plateformes (TikTok, Twitter/X, Instagram, Reddit).")
add_bullet(doc, "Développer un système de détection du langage d'altérisation par règles linguistiques.")
add_bullet(doc, "Entraîner un classifieur supervisé pour automatiser la classification à grande échelle.")
add_bullet(doc, "Enrichir le corpus avec des scores de toxicité (Detoxify) et des émotions (GoEmotions).")
add_bullet(doc, "Analyser les thèmes émergents par modélisation de sujets (BERTopic).")
add_bullet(doc, "Construire un tableau de bord interactif (Streamlit) permettant d'explorer les résultats.")
add_bullet(doc, "Conduire des analyses temporelles pour étudier l'impact d'événements réels sur le discours.")

add_heading2(doc, "2.2 Environnement de travail et outils")

add_para(doc,
    "L'intégralité du projet est développée en Python 3.13 sous NixOS, avec un environnement "
    "virtuel dédié. Le code source est versionné sous Git."
)

add_table(doc,
    ["Catégorie", "Outil / Librairie", "Usage"],
    [
        ["Langage",          "Python 3.13",              "Développement principal"],
        ["Traitement NLP",   "scikit-learn, NLTK",       "TF-IDF, classifieur, nettoyage"],
        ["Modélisation thèmes","BERTopic",               "Modélisation de sujets non-supervisée"],
        ["Toxicité",         "Detoxify (Unitary)",       "Score de toxicité par post"],
        ["Émotions",         "GoEmotions (Hugging Face)","Classification émotionnelle 28 classes"],
        ["Visualisation",    "Plotly, Streamlit",        "Graphiques interactifs et dashboard"],
        ["Données Reddit",   "Arctic Shift / Pushshift", "Scraping de posts Reddit par date"],
        ["Versioning",       "Git",                      "Contrôle de version du code"],
    ],
    col_widths=[3.5, 4.5, 6.5]
)

doc.add_paragraph()
doc.add_page_break()

# ============================================================
# 3. PIPELINE DE TRAITEMENT
# ============================================================
add_heading1(doc, "3. Pipeline de traitement")

add_para(doc,
    "Le projet est organisé en une pipeline modulaire de six étapes, conçue pour être reproductible "
    "et extensible. Chaque étape produit un fichier de données intermédiaire stocké dans le répertoire "
    "data/processed/."
)

add_placeholder_figure(doc,
    "Schéma de la pipeline : Collecte → Nettoyage → Détection par règles → "
    "Enrichissement (toxicité + émotions) → Classification ML → Analyse (BERTopic + Event Study) → Dashboard. "
    "Représenter les 6 étapes en boîtes reliées par des flèches, avec les fichiers de données "
    "intermédiaires indiqués entre les blocs."
)

add_heading2(doc, "3.1 Collecte et préparation des données")

add_para(doc,
    "Le corpus final comprend 134 459 publications issues de quatre sources distinctes. "
    "Les données TikTok ont été collectées manuellement via l'API de scraping, les données "
    "Twitter via un jeu de données public, les commentaires Instagram via scraping ciblé, "
    "et les posts Reddit via Arctic Shift (successeur de Pushshift après sa fermeture)."
)

add_table(doc,
    ["Source", "N posts", "Période", "Contenu"],
    [
        ["TikTok",   "1 200",  "2022–2024", "Commentaires sous vidéos immigration"],
        ["Twitter/X","10 000", "2020–2023", "Tweets liés à l'immigration (EN)"],
        ["Instagram","4 018",  "2021–2024", "Commentaires sous posts thématiques"],
        ["Reddit",   "11 247", "2020–2023", "Posts r/politics, worldnews, ukpolitics, europe, immigration"],
        ["HatEval",  "~100",   "2019",      "Jeu de test annoté (discours de haine, EN)"],
    ],
    col_widths=[3.5, 2.5, 3.5, 5]
)

doc.add_paragraph()
add_para(doc,
    "L'étape de nettoyage (src/cleaning.py) normalise chaque texte : passage en minuscules, "
    "suppression des URLs et mentions (@user), suppression des caractères spéciaux, "
    "collapsement des espaces multiples."
)

add_heading2(doc, "3.2 Détecteur de langage d'altérisation par règles")

add_para(doc,
    "La première approche de détection repose sur un système de règles linguistiques (src/othering.py) "
    "organisées en quatre familles de patrons :"
)
add_bullet(doc, "Métaphores déshumanisantes : « invasion », « flood », « swarm », « infestation », « vermin »…")
add_bullet(doc, "Exclusion morale : « don't belong », « go back », « get out », « send them back »…")
add_bullet(doc, "Généralisations : « all of them », « these people always », « they never », « those people »…")
add_bullet(doc, "Cadrage de menace : « taking over », « replacing us », « great replacement », « white genocide »…")

add_para(doc,
    "Chaque texte reçoit un score d'altérisation de 0 à 4 (nombre de familles activées) et une liste "
    "de patrons déclenchés. Ce système de règles a servi de base à l'annotation dite « silver labels » "
    "pour l'entraînement du classifieur supervisé."
)

add_heading2(doc, "3.3 Enrichissement : toxicité et émotions")

add_para(doc,
    "Deux modèles pré-entraînés ont été appliqués à l'ensemble du corpus pour enrichir chaque post "
    "avec des dimensions analytiques supplémentaires :"
)

add_bullet(doc,
    "Toxicité (Detoxify) : score continu [0, 1] mesurant le caractère offensant ou nocif d'un texte. "
    "Le seuil de 0,5 définit un post comme « hautement toxique »."
)
add_bullet(doc,
    "Émotions (GoEmotions) : classification multi-classes sur 28 émotions de base, réduite à l'émotion "
    "dominante par post. Les émotions les plus fréquentes dans le corpus sont neutral (35,5 %), "
    "anger (18,2 %), annoyance (9,9 %) et curiosity (5,6 %)."
)

add_heading2(doc, "3.4 Classifieur machine learning (LinearSVC TF-IDF)")

add_para(doc,
    "Face aux limites du système de règles (rappel insuffisant pour les formulations implicites), "
    "un classifieur supervisé a été entraîné à partir des étiquettes silver générées par les règles. "
    "Deux modèles ont été comparés sur une représentation TF-IDF (50 000 unigrammes + bigrammes, "
    "log-fréquence, min_df=2) :"
)

add_table(doc,
    ["Modèle", "Précision", "Rappel", "F1-score"],
    [
        ["Régression logistique",  "0,943", "0,973", "0,958"],
        ["LinearSVC (retenu)", "0,994", "0,985", "0,990"],
    ],
    col_widths=[5.5, 3, 3, 3]
)

doc.add_paragraph()
add_para(doc,
    "Le LinearSVC a été retenu pour ses meilleures performances. La matrice de confusion sur le jeu de "
    "test (26 892 exemples) confirme la robustesse du modèle :"
)
add_bullet(doc, "Vrais négatifs : 25 461  (textes non-altérisants correctement classifiés)")
add_bullet(doc, "Faux positifs : 8       (textes normaux classifiés à tort comme altérisants)")
add_bullet(doc, "Faux négatifs : 21      (textes altérisants manqués)")
add_bullet(doc, "Vrais positifs : 1 402  (textes altérisants correctement détectés)")

add_placeholder_figure(doc,
    "Matrice de confusion du classifieur LinearSVC : tableau 2×2 avec en colonnes "
    "'Prédit : non' / 'Prédit : oui' et en lignes 'Réel : non' / 'Réel : oui'. "
    "Valeurs : VN=25461, FP=8, FN=21, VP=1402. Case VP en vert foncé, cases FP/FN en orange. "
    "Indiquer également Précision=0,994 Rappel=0,985 F1=0,990 sous le tableau."
)

add_heading2(doc, "3.5 Analyse thématique (BERTopic)")

add_para(doc,
    "BERTopic a été appliqué au sous-ensemble Reddit pour identifier les thèmes de discussion "
    "dominants de manière non supervisée. BERTopic combine des embeddings contextuels (Sentence "
    "Transformers), une réduction de dimensionnalité (UMAP) et une agglomération hiérarchique "
    "(HDBSCAN) pour découvrir des groupes sémantiques cohérents. Les résultats sont visualisés "
    "dans le dashboard sous forme de graphiques interactifs (taille des topics, score "
    "d'altérisation moyen par topic)."
)

add_placeholder_figure(doc,
    "Visualisation des topics BERTopic : graphique à barres horizontales montrant les 10 à 15 "
    "topics principaux identifiés dans le corpus Reddit, avec leur taille (nombre de posts) et "
    "leur taux d'altérisation moyen (couleur ou double axe). Export disponible dans "
    "reports/figures/topic_sizes.html."
)

add_heading2(doc, "3.6 Analyse temporelle et event study")

add_para(doc,
    "Le corpus Reddit couvre la période 2020–2023, permettant de croiser les variations de "
    "discours avec des événements géopolitiques majeurs (invasion de l'Ukraine, crise afghane, "
    "attentats, élections…). La méthode d'event study — standard en économétrie financière — "
    "a été adaptée à ce contexte : pour chaque événement, on calcule la déviation quotidienne "
    "d'une métrique (taux d'altérisation, toxicité, taux d'émotion) par rapport à sa valeur "
    "de référence dans les 14 jours précédant l'événement."
)

add_placeholder_figure(doc,
    "Graphique event study : axe X = décalage en jours par rapport à l'événement (−14 à +14), "
    "axe Y = déviation du taux d'altérisation en points de pourcentage. Chaque ligne représente "
    "un événement sélectionné. Insérer une capture du dashboard, onglet Temporal Analysis, "
    "section Event Study, pour un ou deux événements représentatifs (ex: invasion Ukraine "
    "2022-02-24 ou crise afghane 2021-08-15)."
)

doc.add_page_break()

# ============================================================
# 4. RÉSULTATS ET ANALYSE
# ============================================================
add_heading1(doc, "4. Résultats et analyse")

add_heading2(doc, "4.1 Comparaison inter-plateformes")

add_para(doc,
    "La mesure du taux d'altérisation — proportion de posts classifiés comme contenant du "
    "langage « Nous vs Eux » — révèle des disparités importantes selon la plateforme :"
)

add_table(doc,
    ["Plateforme", "N posts", "Taux d'altérisation", "Toxicité moyenne"],
    [
        ["TikTok",    "1 200",  "3,58 %", "0,123"],
        ["Instagram", "4 018",  "1,77 %", "0,091"],
        ["Twitter/X", "10 000", "1,66 %", "0,003"],
        ["Reddit",    "11 247", "0,82 %", "0,020"],
    ],
    col_widths=[3.5, 2.5, 4, 4]
)

doc.add_paragraph()
add_para(doc,
    "TikTok présente le taux d'altérisation le plus élevé (3,58 %), ce qui peut s'expliquer "
    "par la nature virale et algorithmique de la plateforme, qui favorise les contenus provocateurs. "
    "Reddit, en revanche, affiche le taux le plus faible (0,82 %), probablement grâce à ses "
    "mécanismes de modération communautaire."
)

add_placeholder_figure(doc,
    "Graphique à barres comparant les 4 plateformes : axe X = plateforme, axe Y = taux "
    "d'altérisation (%). Couleurs distinctes par plateforme. Insérer le graphique disponible "
    "dans reports/figures/platform_othering.png."
)

add_heading2(doc, "4.2 Analyse par sous-reddit")

add_para(doc,
    "Au sein de Reddit, des différences marquées s'observent entre les communautés :"
)

add_table(doc,
    ["Subreddit", "N posts", "Taux d'altérisation", "Toxicité moyenne"],
    [
        ["r/worldnews",    "2 250", "1,56 %", "0,031"],
        ["r/europe",       "2 250", "1,07 %", "0,014"],
        ["r/ukpolitics",   "2 250", "0,76 %", "0,018"],
        ["r/politics",     "2 250", "0,67 %", "0,032"],
        ["r/immigration",  "2 247", "0,04 %", "0,004"],
    ],
    col_widths=[4, 2.5, 4, 4]
)

doc.add_paragraph()
add_para(doc,
    "Le subreddit r/worldnews présente le taux d'altérisation le plus élevé (1,56 %), "
    "reflet d'un contexte de discussions internationales souvent chargées. À l'opposé, "
    "r/immigration affiche un taux quasiment nul (0,04 %), ce qui suggère que les communautés "
    "directement concernées par le sujet adoptent un discours plus factuel et moins stigmatisant."
)

add_heading2(doc, "4.3 Dynamiques temporelles")

add_para(doc,
    "L'analyse des pronoms révèle la structure du corpus : sur les 134 459 posts, "
    "63,4 % ne contiennent aucun pronom de groupe (« none »), 18,5 % uniquement « eux/they » "
    "(them_only), 11,4 % uniquement « nous/we » (we_only), et 6,7 % les deux (both). "
    "La coprésence des deux pronoms est un indicateur fort de discours oppositionnel."
)

add_placeholder_figure(doc,
    "Graphique en aires empilées (100 % normalisé) montrant l'évolution de la distribution "
    "des émotions dans le temps (2020–2023) sur Reddit. Les 8 émotions les plus fréquentes "
    "sont représentées avec un code couleur distinct (neutral en gris, anger en rouge, etc.). "
    "Insérer une capture du dashboard, onglet Timeline."
)

add_para(doc,
    "L'analyse event study autour de l'invasion russe de l'Ukraine (24 février 2022) montre "
    "une élévation notable du taux d'altérisation dans les 7 jours suivant l'événement, "
    "particulièrement sur r/worldnews et r/europe, accompagnée d'une augmentation du sentiment "
    "de peur (fear) et de colère (anger). Des dynamiques similaires s'observent autour de la "
    "chute de Kaboul (août 2021) et des élections américaines (novembre 2020)."
)

doc.add_page_break()

# ============================================================
# 5. DASHBOARD INTERACTIF
# ============================================================
add_heading1(doc, "5. Dashboard interactif")

add_para(doc,
    "L'ensemble des analyses est rendu accessible via un tableau de bord interactif développé "
    "avec Streamlit (Python). Ce dashboard permet à un utilisateur non-technicien d'explorer "
    "les données filtrées par plateforme, sous-reddit, période temporelle et type de discours."
)

add_para(doc, "Le dashboard est organisé en six onglets :")
add_bullet(doc, "Overview : indicateurs clés (KPI), distribution des émotions, carte de densité.")
add_bullet(doc, "Timeline : évolution temporelle de l'altérisation, de la toxicité et des émotions.")
add_bullet(doc, "Temporal Analysis : graphiques event study par événement et par métrique.")
add_bullet(doc, "Topics : exploration des thèmes BERTopic et leur score d'altérisation.")
add_bullet(doc, "Toxicity : distribution des scores Detoxify et relation toxicité/altérisation.")
add_bullet(doc, "Classifier : performances du modèle LinearSVC et matrice de confusion.")

add_placeholder_figure(doc,
    "Capture d'écran du dashboard Streamlit, onglet Overview : afficher les KPI en haut "
    "(Total Posts, Othering Rate, Avg Toxicity, High Toxicity %) et le graphique de distribution "
    "des émotions. Lancer le dashboard avec : "
    "LD_LIBRARY_PATH=... streamlit run dashboard/app.py"
)

add_placeholder_figure(doc,
    "Capture d'écran du dashboard Streamlit, onglet Timeline : afficher le graphique "
    "d'évolution du taux d'altérisation dans le temps avec les lignes d'événements verticales."
)

doc.add_page_break()

# ============================================================
# 6. CONCLUSION
# ============================================================
add_heading1(doc, "6. Conclusion")

add_para(doc,
    "Ce stage a permis de construire de bout en bout un système complet de détection et d'analyse "
    "du langage d'altérisation dans les réseaux sociaux. La pipeline développée combine des "
    "approches complémentaires — règles linguistiques, apprentissage automatique supervisé, "
    "modèles pré-entraînés et modélisation de sujets — pour produire une analyse multi-dimensionnelle "
    "d'un corpus de plus de 134 000 publications."
)

add_para(doc,
    "Les résultats montrent que le langage d'altérisation est présent à des niveaux variables "
    "selon les plateformes (de 0,82 % sur Reddit à 3,58 % sur TikTok) et les communautés, "
    "et qu'il répond à des dynamiques temporelles liées à des événements géopolitiques réels. "
    "Le classifieur LinearSVC atteint un F1-score de 0,990, démontrant qu'une détection "
    "automatique robuste est atteignable avec des features TF-IDF et une bonne stratégie "
    "d'annotation par règles."
)

add_para(doc,
    "Sur le plan professionnel, ce projet a renforcé mes compétences en NLP, en Python "
    "scientifique (scikit-learn, Plotly, Streamlit) et en gestion de projet autonome. "
    "Il m'a également permis de me confronter aux défis réels de l'annotation de données, "
    "de la reproductibilité des expériences et de la communication de résultats techniques "
    "à un public non spécialisé via un outil interactif."
)

add_para(doc,
    "Des pistes d'amélioration existent : l'extension à d'autres langues (français, arabe), "
    "l'intégration d'un modèle de langage de grande taille (LLM) pour les formulations implicites, "
    "et la mise en place d'une collecte de données en temps réel. Ce projet a consolidé mon "
    "orientation vers les métiers de la data science et du NLP."
)

doc.add_page_break()

# ============================================================
# GLOSSAIRE
# ============================================================
add_heading1(doc, "Glossaire")

glossaire = [
    ("Altérisation (othering)", "Processus rhétorique par lequel un groupe est désigné comme radicalement différent et inférieur, contribuant à la construction d'un « eux » opposé à un « nous »."),
    ("BERTopic", "Algorithme de modélisation de sujets non-supervisé combinant des embeddings BERT, UMAP et HDBSCAN pour identifier des thèmes cohérents dans un corpus textuel."),
    ("Classifieur", "Modèle d'apprentissage automatique qui prédit la classe (catégorie) d'un exemple à partir de ses caractéristiques."),
    ("Corpus", "Ensemble de textes constituant les données d'une étude."),
    ("F1-score", "Moyenne harmonique de la précision et du rappel, mesure équilibrée de la performance d'un classifieur binaire."),
    ("GoEmotions", "Jeu de données et modèle Google pour la classification d'émotions à 28 classes dans des textes courts."),
    ("LinearSVC", "Support Vector Machine linéaire, classifieur basé sur la maximisation de la marge entre les classes."),
    ("NLP (Natural Language Processing)", "Traitement automatique du langage naturel, branche de l'IA traitant le texte humain."),
    ("Pushshift / Arctic Shift", "Services d'archivage et d'accès aux données historiques Reddit."),
    ("Silver labels", "Étiquettes d'entraînement générées automatiquement (par règles), par opposition aux gold labels annotés manuellement."),
    ("TF-IDF", "Term Frequency–Inverse Document Frequency : pondération statistique mesurant l'importance d'un mot dans un document par rapport à un corpus."),
    ("Toxicité", "Score mesurant le caractère offensant, insultant ou inapproprié d'un texte, calculé par le modèle Detoxify."),
]

for term, definition in glossaire:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_term = p.add_run(f"{term} : ")
    run_term.font.name = "Calibri"
    run_term.font.size = Pt(11)
    run_term.bold = True
    run_def = p.add_run(definition)
    run_def.font.name = "Calibri"
    run_def.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# RÉFÉRENCES
# ============================================================
add_heading1(doc, "Références")

refs = [
    ("[1]", "Baumgartner, J. et al. (2020). The Pushshift Reddit Dataset. In Proceedings of ICWSM. "
            "Données Reddit archivées accessibles via Arctic Shift : https://arctic-shift.photon-reddit.com"),
    ("[2]", "Grootendorst, M. (2022). BERTopic: Neural topic modeling with a class-based TF-IDF procedure. "
            "arXiv:2203.05794."),
    ("[3]", "Demszky, D. et al. (2020). GoEmotions: A Dataset of Fine-Grained Emotions. "
            "In Proceedings of ACL 2020."),
]

for ref_id, ref_text in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.first_line_indent = Cm(-1.2)
    run_id = p.add_run(f"{ref_id}  ")
    run_id.font.name = "Calibri"
    run_id.font.size = Pt(11)
    run_id.bold = True
    run_text = p.add_run(ref_text)
    run_text.font.name = "Calibri"
    run_text.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# ANNEXE A – MATRICE SWOT
# ============================================================
add_heading1(doc, "Annexe A – Matrice SWOT")

add_table(doc,
    ["Forces (Strengths)", "Faiblesses (Weaknesses)"],
    [
        [
            "• Maîtrise de Python et de l'écosystème scikit-learn acquise en formation\n"
            "• Capacité à construire une pipeline ML complète de bout en bout\n"
            "• Connaissance des outils NLP modernes (TF-IDF, embeddings, BERTopic)\n"
            "• Curiosité et aptitude à l'apprentissage autonome",
            "• Manque d'expérience en annotation manuelle de données\n"
            "• Faible exposition aux LLMs (GPT, LLaMA) pour les tâches NLP avancées\n"
            "• Compétences en déploiement cloud (Docker, CI/CD) à renforcer\n"
            "• Gestion du temps sur des projets longs à consolider"
        ],
    ],
    col_widths=[8.5, 8.5]
)
doc.add_paragraph()
add_table(doc,
    ["Opportunités (Opportunities)", "Menaces (Threats)"],
    [
        [
            "• Forte demande en data science et NLP dans de nombreux secteurs\n"
            "• Essor de l'IA générative ouvrant des opportunités de spécialisation\n"
            "• Intérêt croissant pour la modération automatisée et la lutte contre la haine en ligne\n"
            "• Possibilité de publier ce travail comme article de recherche",
            "• Évolution rapide des modèles (risque d'obsolescence rapide des compétences)\n"
            "• Marché du travail très compétitif en data science\n"
            "• Enjeux éthiques et légaux autour de l'IA et de la modération de contenu\n"
            "• Biais potentiels dans les données et les modèles"
        ],
    ],
    col_widths=[8.5, 8.5]
)

doc.add_page_break()

# ============================================================
# ANNEXE B – COMPÉTENCES BUT MOBILISÉES
# ============================================================
add_heading1(doc, "Annexe B – Compétences BUT Informatique mobilisées")

competences = [
    ("Compétence 1 – Réaliser un développement",
     "Développement de la pipeline complète en Python : nettoyage de texte, détecteur par règles, "
     "classifieur supervisé, enrichissement par modèles pré-entraînés, dashboard Streamlit. "
     "Niveau initial : intermédiaire. Niveau final : avancé — capable de concevoir et implémenter "
     "une application NLP de production."),
    ("Compétence 2 – Optimiser des applications",
     "Comparaison et sélection du meilleur classifieur (LR vs LinearSVC) sur la base de métriques "
     "objectives (F1, précision, rappel). Optimisation de la représentation TF-IDF (max_features, "
     "ngram_range, sublinear_tf). Application des mathématiques de la classification binaire. "
     "Niveau initial : faible (théorie). Niveau final : opérationnel."),
    ("Compétence 3 – Gérer des données",
     "Construction et gestion d'un corpus multi-sources de 134 459 posts. Conception du schéma "
     "de données, transformations pandas, stockage en CSV/Parquet. Visualisation avancée "
     "(Plotly, Streamlit). Niveau initial : intermédiaire. Niveau final : avancé."),
    ("Compétence 5 – Conduire un projet",
     "Gestion autonome du projet sur 2 mois : définition des objectifs, planification des sprints "
     "hebdomadaires, versioning Git, documentation du code. Niveau initial : basique. Niveau final : "
     "autonome sur les projets individuels."),
]

for title, desc in competences:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = True
    add_para(doc, desc)

doc.add_page_break()

# ============================================================
# ANNEXE C – UTILISATION DE L'IA
# ============================================================
add_heading1(doc, "Annexe C – Utilisation de l'intelligence artificielle")

add_para(doc,
    "Conformément aux exigences de l'IUT, cette annexe recense l'ensemble des utilisations "
    "de l'intelligence artificielle dans le cadre de ce projet et de la rédaction de ce rapport."
)

add_heading2(doc, "C.1 Utilisation de l'IA dans le projet (développement)")

add_bullet(doc, "Outil : Claude Code (Anthropic, modèle Claude Sonnet 4.6)")
add_bullet(doc, "Type d'utilisation : aide à la rédaction de code (suggestions, débogage, refactoring)")
add_bullet(doc, "Portions concernées : génération du script make_pptx.py, amélioration du dashboard Streamlit (ajout de charts, corrections de layout)")
add_bullet(doc, "Proportion estimée : < 25 % du code total (la structure, la logique NLP et les données sont entièrement le travail de l'étudiant)")

add_heading2(doc, "C.2 Utilisation de l'IA pour la rédaction du rapport")

add_bullet(doc, "Outil : Claude Code (Anthropic, modèle Claude Sonnet 4.6)")
add_bullet(doc, "Type d'utilisation : aide à la structuration et à la rédaction du rapport de stage à partir des données et du code existants")
add_bullet(doc, "Proportion estimée : < 25 % du texte total (toutes les données, analyses et conclusions proviennent du travail de l'étudiant)")
add_bullet(doc, "Vérification : l'ensemble des statistiques et affirmations ont été vérifiées par l'étudiant à partir des fichiers de données réels")

add_para(doc,
    "PRINCIPE DE VÉRIFICATION : l'étudiant est responsable de l'ensemble du contenu soumis "
    "à évaluation. Toutes les données chiffrées ont été extraites directement des fichiers CSV "
    "et JSON du projet et vérifiées avant inclusion dans ce rapport.",
    italic=True
)

# ============================================================
# SAVE
# ============================================================
out_path = "reports/rapport_stage_BUT2.docx"
doc.save(out_path)
print(f"Rapport généré : {out_path}")

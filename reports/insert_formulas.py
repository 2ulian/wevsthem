"""
Insère les formules mathématiques dans le rapport existant,
en recherchant les paragraphes d'ancrage par leur contenu textuel.
Ne régénère pas le document — modifie uniquement le DOCX en place.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from lxml import etree

DOC_PATH = "reports/rapport_stage_BUT2.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def make_formula_para(doc, label, formula, note=None):
    """
    Returns a list of paragraphs representing a display formula block.
    label   : e.g. "TF-IDF(t, d)"
    formula : the formula string
    note    : optional explanation line
    """
    paragraphs = []

    # Label + formula on one line, centered, box-like
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(2)
    p.paragraph_format.right_indent = Cm(2)

    # Gray shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)

    if label:
        run_label = p.add_run(f"{label}  =  ")
        run_label.font.name = "Calibri"
        run_label.font.size = Pt(11)
        run_label.font.bold = True
        run_label.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    run_formula = p.add_run(formula)
    run_formula.font.name = "Consolas"
    run_formula.font.size = Pt(11)
    run_formula.font.bold = False
    run_formula.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    paragraphs.append(p)

    if note:
        p2 = doc.add_paragraph()
        p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(8)
        p2.paragraph_format.left_indent  = Cm(2)
        p2.paragraph_format.right_indent = Cm(2)
        run_note = p2.add_run(note)
        run_note.font.name = "Calibri"
        run_note.font.size = Pt(9)
        run_note.font.italic = True
        run_note.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        paragraphs.append(p2)

    return paragraphs


def find_para_index(doc, search_text):
    """Return index of first paragraph whose text starts with search_text."""
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(search_text):
            return i
    return None


def insert_paragraphs_after(doc, anchor_idx, new_paragraphs):
    """
    Insert new_paragraphs into doc.element.body immediately after the
    paragraph at doc.paragraphs[anchor_idx].
    """
    # Get the XML element of the anchor paragraph
    anchor_elem = doc.paragraphs[anchor_idx]._p

    # Insert each new paragraph after the previous insertion point
    prev = anchor_elem
    for p in new_paragraphs:
        # Remove from wherever add_paragraph() placed it (end of body)
        p._p.getparent().remove(p._p)
        # Re-insert after prev
        prev.addnext(p._p)
        prev = p._p


# ---------------------------------------------------------------------------
# Load document
# ---------------------------------------------------------------------------

doc = Document(DOC_PATH)

insertions_done = []

# ============================================================
# FORMULE 1 — Score d'altérisation par règles  (section 3.2)
# Après : "Chaque texte reçoit un score d'altérisation de 0 à 4"
# ============================================================
anchor = find_para_index(doc, "Chaque texte reçoit un score d'altérisation")
if anchor is not None:
    paras = make_formula_para(
        doc,
        label="othering_score(t)",
        formula="| { familles ∈ {déhumanisant, exclusion, généralisation, menace} : famille activée dans t } |",
        note="Score entier ∈ {0, 1, 2, 3, 4} — nombre de familles de patrons déclenchées dans le texte t"
    )
    insert_paragraphs_after(doc, anchor, paras)
    insertions_done.append("othering_score (3.2)")
else:
    print("WARNING: anchor for othering_score not found")

# ============================================================
# FORMULE 2 — TF-IDF  (section 3.4)
# Après : "Face aux limites du système de règles"
# ============================================================
anchor = find_para_index(doc, "Face aux limites du système de règles")
if anchor is not None:
    paras = make_formula_para(
        doc,
        label="TF-IDF(t, d)",
        formula="log(1 + tf(t, d))  ×  log( N / df(t) )",
        note="tf(t,d) = fréquence du terme t dans le document d  |  N = taille du corpus  |  df(t) = nombre de documents contenant t"
    )
    insert_paragraphs_after(doc, anchor, paras)
    insertions_done.append("TF-IDF (3.4)")
else:
    print("WARNING: anchor for TF-IDF not found")

# ============================================================
# FORMULE 3 — Précision, Rappel, F1  (section 3.4)
# Après : "Le LinearSVC a été retenu pour ses meilleures performances"
# ============================================================
anchor = find_para_index(doc, "Le LinearSVC a été retenu pour ses meilleures performances")
if anchor is not None:
    paras = []
    paras += make_formula_para(
        doc,
        label="Précision",
        formula="TP / (TP + FP)  =  1402 / (1402 + 8)  =  0,994",
        note="TP = vrais positifs  |  FP = faux positifs"
    )
    paras += make_formula_para(
        doc,
        label="Rappel",
        formula="TP / (TP + FN)  =  1402 / (1402 + 21)  =  0,985",
        note="FN = faux négatifs"
    )
    paras += make_formula_para(
        doc,
        label="F1-score",
        formula="2 × (Précision × Rappel) / (Précision + Rappel)  =  0,990",
        note="Moyenne harmonique — pénalise les déséquilibres entre précision et rappel"
    )
    insert_paragraphs_after(doc, anchor, paras)
    insertions_done.append("Précision / Rappel / F1 (3.4)")
else:
    print("WARNING: anchor for P/R/F1 not found")

# ============================================================
# FORMULE 4 — Déviation event study  (section 3.6)
# Après : "Le corpus Reddit couvre la période 2020–2023"
# ============================================================
anchor = find_para_index(doc, "Le corpus Reddit couvre la période 2020")
if anchor is not None:
    paras = make_formula_para(
        doc,
        label="déviation(t)",
        formula="métrique(t)  −  moyenne( métrique(τ) )  pour τ ∈ [t_evt − 14j, t_evt[",
        note="t = jour courant  |  t_evt = date de l'événement  |  baseline = fenêtre pré-événement de 14 jours"
    )
    insert_paragraphs_after(doc, anchor, paras)
    insertions_done.append("déviation event study (3.6)")
else:
    print("WARNING: anchor for event study not found")

# ============================================================
# FORMULE 5 — Taux d'altérisation  (section 4.1)
# Après : "La mesure du taux d'altérisation — proportion de posts"
# ============================================================
anchor = find_para_index(doc, "La mesure du taux d'altérisation")
if anchor is not None:
    paras = make_formula_para(
        doc,
        label="othering_rate",
        formula="( Σ othering_predicted(i)  /  N )  ×  100  (en %)",
        note="othering_predicted(i) ∈ {0, 1}  |  N = nombre total de posts dans le périmètre considéré"
    )
    insert_paragraphs_after(doc, anchor, paras)
    insertions_done.append("othering_rate (4.1)")
else:
    print("WARNING: anchor for othering_rate not found")

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(DOC_PATH)
print(f"\nFormules insérées ({len(insertions_done)}) :")
for x in insertions_done:
    print(f"  ✓ {x}")
print(f"\nFichier mis à jour : {DOC_PATH}")
